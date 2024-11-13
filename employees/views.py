from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from django.conf import settings
from django.contrib.auth import authenticate, login, logout
from django.utils.dateformat import DateFormat
from django.contrib.auth.decorators import login_required
from django.core.exceptions import PermissionDenied
from django.utils import timezone
from django.contrib import messages
from django.urls import reverse
from django.contrib.auth.models import User
from .models import Employee, Contract
from .forms import EmployeeForm
from docx import Document
from pathlib import Path
import os
import re

# Definiujemy ścieżkę do szablonów
templates_directory = Path(settings.MEDIA_ROOT) / 'documents'

def index(request):
    return redirect('employees:login')

def user_login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('employees:dashboard')
        else:
            return render(request, 'login.html', {'error': 'Nieprawidłowe dane logowania'})
    return render(request, 'login.html')

def user_logout(request):
    logout(request)
    return redirect('employees:login')

def get_placeholders(doc_path):
    """Extract placeholders from a .docx file enclosed in {}."""
    placeholders = set()
    doc = Document(doc_path)
    pattern = re.compile(r"\{(\w+)\}")  # Matches patterns like {name}
    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        placeholders.update(matches)
    return list(placeholders)

@login_required
def dashboard(request):
    # Ścieżka do katalogu `media/documents` dla szablonów
    templates_directory = Path(settings.MEDIA_ROOT) / 'documents'

    # Sprawdzenie, czy użytkownik jest w grupie Admin
    if request.user.groups.filter(name='Admin').exists():
        employees = Employee.objects.all()
        contracts = Contract.objects.all()
        
        # Sprawdzenie, czy katalog z szablonami istnieje, i pobranie plików .docx
        if templates_directory.exists():
            templates = [t.name for t in templates_directory.glob("*.docx") if t.is_file()]
        else:
            templates = []
            print("Błąd: Katalog `media/documents` nie istnieje.")  # Debugowanie: upewnij się, że katalog istnieje
        
        return render(request, 'employees/admin_dashboard.html', {
            'employees': employees,
            'contracts': contracts,
            'templates': templates,
        })

    # Sprawdzenie, czy użytkownik jest w grupie Pracownik
    elif request.user.groups.filter(name='Pracownik').exists():
        employee = get_object_or_404(Employee, user=request.user)
        contracts = Contract.objects.filter(employee=employee)
        return render(request, 'employees/employee_dashboard.html', {
            'employee': employee,
            'contracts': contracts
        })

    # Jeśli użytkownik nie należy do żadnej z grup
    else:
        return redirect('employees:login')

@login_required
def employee_list(request):
    if request.user.groups.filter(name='Admin').exists():
        employees = Employee.objects.all()
        return render(request, 'employees/employee_list.html', {'employees': employees})
    else:
        raise PermissionDenied

@login_required
def employee_create(request):
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        email = request.POST.get('email')

        # Generate username by combining first and last names
        username = (first_name + last_name).lower()
        
        # Create the User object
        user = User.objects.create_user(username=username, password="123", email=email)
        
        # Create the Employee profile associated with the new User
        Employee.objects.create(user=user, first_name=first_name, last_name=last_name, email=email)

        return redirect('employees:dashboard')
    return render(request, 'employees/employee_form.html')

@login_required
def employee_detail(request, pk):
    employee = get_object_or_404(Employee, pk=pk)
    is_admin = request.user.groups.filter(name='Admin').exists()

    # Fetch available templates
    templates = [t.name for t in templates_directory.glob("*.docx") if t.is_file()]
    
    # Fetch all contracts for the employee
    contracts = Contract.objects.filter(employee=employee).order_by('-generated_date')

    # Check if the user has permission to view this page
    if not is_admin and employee.user != request.user:
        raise PermissionDenied

    return render(request, 'employees/employee_detail.html', {
        'employee': employee,
        'is_admin': is_admin,
        'templates': templates,
        'contracts': contracts,
    })

@login_required
def generate_contract(request, pk):
    if request.method == 'POST':
        employee = get_object_or_404(Employee, pk=pk)
        template_name = request.POST.get('template')
        document_name = request.POST.get('document_name', f"Contract_{employee.first_name}_{employee.last_name}.docx")
        
        if not document_name.endswith('.docx'):
            document_name += '.docx'

        # Ścieżka do szablonu
        template_path = Path(settings.MEDIA_ROOT) / 'documents' / template_name
        if not template_path.exists():
            return JsonResponse({'status': 'error', 'message': 'Szablon nie istnieje.'}, status=400)

        # Tworzenie i wypełnianie dokumentu
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for placeholder, value in request.POST.items():
                if placeholder.startswith("placeholder_"):
                    placeholder_name = placeholder.replace("placeholder_", "")
                    paragraph.text = paragraph.text.replace(f'{{{placeholder_name}}}', value)

        # Zapisywanie dokumentu
        employee_dir = Path(settings.MEDIA_ROOT) / 'contracts' / f"{employee.first_name}_{employee.last_name}"
        employee_dir.mkdir(parents=True, exist_ok=True)
        contract_path = employee_dir / document_name
        doc.save(contract_path)

        # Zapis w bazie danych
        contract = Contract.objects.create(
            employee=employee,
            document_path=str(contract_path.relative_to(settings.MEDIA_ROOT)),
            document_name=document_name,
            generated_date=timezone.now()
        )

        return JsonResponse({
            'status': 'success',
            'contract': {
                'document_name': contract.document_name,
                'generated_date': contract.generated_date.strftime('%d/%m/%Y'),
                'download_url': f"{settings.MEDIA_URL}{contract.document_path}"
            }
        }, status=200)
    else:
        return JsonResponse({'status': 'error', 'message': 'Nieprawidłowa metoda.'}, status=405)
    
@login_required
def fetch_placeholders(request):
    template_name = request.GET.get('template')
    template_path = Path(settings.MEDIA_ROOT) / 'documents' / template_name
    placeholders = []

    if template_path.exists():
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            matches = re.findall(r'\{(\w+)\}', paragraph.text)
            placeholders.extend(matches)

        # Usuń duplikaty, jeśli istnieją
        placeholders = list(set(placeholders))
        return JsonResponse({'status': 'success', 'placeholders': placeholders})
    else:
        return JsonResponse({'status': 'error', 'message': 'Szablon nie znaleziony.'}, status=404)

@login_required
def select_template(request, pk):
    employee = get_object_or_404(Employee, pk=pk)
    # Fetch available templates
    templates = [t.name for t in templates_directory.glob("*.docx") if t.is_file()]
    
    # Render the template selection page
    return render(request, 'select_template.html', {
        'employee': employee,
        'templates': templates
    })

@login_required
def delete_contract(request, pk):
    if request.user.groups.filter(name='Admin').exists():
        contract = get_object_or_404(Contract, pk=pk)
        contract_path = Path(settings.MEDIA_ROOT) / contract.document_path
        if contract_path.exists():
            contract_path.unlink()
        contract.delete()
        messages.success(request, f'Umowa {contract.document_name} została pomyślnie usunięta.')
        return redirect('employees:dashboard')
    else:
        raise PermissionDenied

@login_required
def generate_contract(request, pk):
    if request.method == 'POST':
        employee = get_object_or_404(Employee, pk=pk)
        template_name = request.POST.get('template')
        document_name = request.POST.get('document_name', f"Contract_{employee.first_name}_{employee.last_name}.docx")
        
        if not document_name.endswith('.docx'):
            document_name += '.docx'

        # Ścieżka do szablonu
        template_path = Path(settings.MEDIA_ROOT) / 'documents' / template_name
        if not template_path.exists():
            return JsonResponse({'status': 'error', 'message': 'Szablon nie istnieje.'}, status=400)

        # Tworzenie i wypełnianie dokumentu
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for placeholder, value in request.POST.items():
                if placeholder.startswith("placeholder_"):
                    placeholder_name = placeholder.replace("placeholder_", "")
                    paragraph.text = paragraph.text.replace(f'{{{placeholder_name}}}', value)

        # Zapisywanie dokumentu
        employee_dir = Path(settings.MEDIA_ROOT) / 'contracts' / f"{employee.first_name}_{employee.last_name}"
        employee_dir.mkdir(parents=True, exist_ok=True)
        contract_path = employee_dir / document_name
        doc.save(contract_path)

        # Zapis w bazie danych
        contract = Contract.objects.create(
            employee=employee,
            document_path=str(contract_path.relative_to(settings.MEDIA_ROOT)),
            document_name=document_name,
            generated_date=timezone.now()
        )

        # Odpowiedź JSON zawierająca szczegóły umowy
        return JsonResponse({
            'status': 'success',
            'contract': {
                'document_name': contract.document_name,
                'generated_date': contract.generated_date.strftime('%d/%m/%Y'),
                'download_url': f"{settings.MEDIA_URL}{contract.document_path}",
                'employee_name': f"{employee.first_name} {employee.last_name}"
            }
        }, status=200)
    else:
        return JsonResponse({'status': 'error', 'message': 'Nieprawidłowa metoda.'}, status=405)


@login_required
def get_employee_data(request):
    employee_id = request.GET.get('id')
    if not employee_id:
        return JsonResponse({'status': 'error', 'message': 'Brak ID pracownika.'}, status=400)

    try:
        employee = Employee.objects.get(pk=employee_id)
        return JsonResponse({
            'status': 'success',
            'employee': {
                'first_name': employee.first_name,
                'last_name': employee.last_name,
                'position': employee.position,
                'hire_date': employee.hire_date.strftime('%Y-%m-%d')
            }
        }, status=200)
    except Employee.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Pracownik nie istnieje.'}, status=404)